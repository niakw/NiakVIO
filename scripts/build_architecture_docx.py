from pathlib import Path
import argparse
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def build(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.58)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(3.4)
    normal.paragraph_format.line_spacing = 1.06

    for name, size, color, space_before, space_after in [
        ("Title", 24, "17213D", 0, 10),
        ("Heading 1", 15.2, "17213D", 12, 5),
        ("Heading 2", 11.5, "334A70", 8, 3),
        ("Heading 3", 10.3, "334A70", 6, 2),
    ]:
        style = styles[name]
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.keep_with_next = True

    if "QuoteNote" not in styles:
        quote = styles.add_style("QuoteNote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["QuoteNote"]
    quote.font.name = "Liberation Sans"
    quote.font.size = Pt(9)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(80, 92, 118)
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.08)
    quote.paragraph_format.space_after = Pt(7)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Liberation Mono"
    code.font.size = Pt(7.8)
    code.font.color.rgb = RGBColor(37, 52, 91)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    header = sec.header.paragraphs[0]
    header.text = "NiakVIO  ·  Provider v3 Architecture"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Liberation Sans"
        run.font.size = Pt(7.4)
        run.font.color.rgb = RGBColor(108, 118, 140)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NiakVIO  ·  Architecture  ·  ")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(7.2)
    run.font.color.rgb = RGBColor(108, 118, 140)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    inline_re = re.compile(r"(\*\*.*?\*\*|`[^`]+`)")

    def add_inline(paragraph, value: str) -> None:
        pos = 0
        for match in inline_re.finditer(value):
            if match.start() > pos:
                plain = paragraph.add_run(value[pos:match.start()])
                plain.font.name = "Liberation Sans"
            token = match.group(0)
            if token.startswith("**"):
                formatted = paragraph.add_run(token[2:-2])
                formatted.bold = True
                formatted.font.name = "Liberation Sans"
            else:
                formatted = paragraph.add_run(token[1:-1])
                formatted.font.name = "Liberation Mono"
                formatted.font.size = Pt(8.2)
                formatted.font.color.rgb = RGBColor(45, 70, 112)
            pos = match.end()
        if pos < len(value):
            plain = paragraph.add_run(value[pos:])
            plain.font.name = "Liberation Sans"

    def shade_paragraph(paragraph, fill="F4F7FB", border="D7DFEE") -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), fill)
        p_pr.append(shade)
        borders = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            element = OxmlElement("w:" + side)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "4")
            element.set(qn("w:color"), border)
            borders.append(element)
        p_pr.append(borders)

    def bullet(paragraph, level=0) -> None:
        paragraph.style = styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.add_run("• ").bold = True

    def numbered(paragraph, number: str) -> None:
        paragraph.style = styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        prefix = paragraph.add_run(f"{number}. ")
        prefix.bold = True

    in_code = False
    code_lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                paragraph = doc.add_paragraph(style="CodeBlock")
                paragraph.add_run("\n".join(code_lines))
                shade_paragraph(paragraph)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run(line[2:].strip())
            rule = doc.add_paragraph()
            rule.paragraph_format.space_after = Pt(5)
            accent = rule.add_run("━" * 44)
            accent.font.size = Pt(7)
            accent.font.color.rgb = RGBColor(86, 132, 190)
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph(style="QuoteNote")
            add_inline(paragraph, line[2:].strip())
            shade_paragraph(paragraph, "F7F9FC", "DCE4F0")
            continue
        match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if match:
            paragraph = doc.add_paragraph()
            numbered(paragraph, match.group(1))
            add_inline(paragraph, match.group(2))
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph()
            bullet(paragraph)
            add_inline(paragraph, line[2:].strip())
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, line.strip())

    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        p_pr.append(OxmlElement("w:widowControl"))
        if paragraph.text.strip().endswith(":") or paragraph.text.strip().endswith(" :"):
            paragraph.paragraph_format.keep_with_next = True

    doc.core_properties.title = "Architecture NiakVIO — Provider v3"
    doc.core_properties.subject = "Normative Provider v3 architecture and release contracts"
    doc.core_properties.author = "NiakVIO"
    doc.core_properties.keywords = "NiakVIO, Nuvio, Provider v3, architecture, release, Labs"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"{output} {output.stat().st_size}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Regenerate ARCHITECTURE.docx from ARCHITECTURE.md")
    parser.add_argument("--source", type=Path, default=Path("ARCHITECTURE.md"))
    parser.add_argument("--output", type=Path, default=Path("ARCHITECTURE.docx"))
    args = parser.parse_args()
    build(args.source, args.output)


if __name__ == "__main__":
    main()
